"""组卷服务：按题型/难度/数量从标准题库抽题，生成 AB 卷 docx（题序+选项乱序）+ 答题卡 + 答案。

样式参照 templates/ 下的真实试卷与答题卡 PDF：
- 试卷：封面页（标题、卷别、考场、姓名/员工编号/座位号/考场号、警示语）→ 分数汇总表 → 分节题目
- AB 卷：同一套题，节内题目顺序与选择题选项顺序均乱序
- 答题卡：选择/判断题为 题号/答案 两行×每组10题 的表格；填空题留空线；简答题留答题区
"""
import random
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from sqlalchemy.orm import Session

from ..config import EXPORT_DIR
from ..models import Question

TYPE_ORDER = ["single", "multiple", "judge", "fill_blank", "short_answer", "essay"]
TYPE_LABELS = {
    "single": "单选题", "multiple": "多选题", "judge": "判断题",
    "fill_blank": "填空题", "short_answer": "简答题", "essay": "论述题",
}
CN_NUMS = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
CHOICE_TYPES = ("single", "multiple", "judge")


# ---------- 抽题 ----------

DIFF_LABELS = {"easy": "简单", "medium": "中等", "hard": "困难"}


def _apply_filters(query, q_type: str, category: str, difficulty: str):
    query = query.filter(Question.status == "approved", Question.q_type == q_type)
    if category and category != "any":
        query = query.filter(Question.category == category)
    if difficulty and difficulty != "any":
        query = query.filter(Question.difficulty == difficulty)
    return query


def select_questions(db: Session, criteria: list[dict]) -> dict[str, list[tuple[Question, int]]]:
    """criteria: [{q_type, category('any'|大类), difficulty('any'|easy|medium|hard), count, score}]
    按条件随机抽题，返回 {q_type: [(Question, score), ...]}，不足则抛错。"""
    picked: dict[str, list[tuple[Question, int]]] = {}
    used_ids: set[int] = set()
    for c in criteria:
        q_type, count, score = c["q_type"], int(c["count"]), int(c.get("score", 1))
        if count <= 0:
            continue
        category = c.get("category", "any")
        difficulty = c.get("difficulty", "any")
        pool = [q for q in _apply_filters(db.query(Question), q_type, category, difficulty).all()
                if q.id not in used_ids]
        if len(pool) < count:
            cat_label = "不限大类" if category in ("any", "") else category
            raise ValueError(
                f"题库不足：{TYPE_LABELS[q_type]}（{cat_label}/{DIFF_LABELS.get(difficulty, '不限难度')}）"
                f"需要 {count} 道，可用仅 {len(pool)} 道，请先补充审核通过的题目")
        chosen = random.sample(pool, count)
        used_ids.update(q.id for q in chosen)
        picked.setdefault(q_type, []).extend((q, score) for q in chosen)
    if not picked:
        raise ValueError("未指定任何抽题条件")
    return picked


def find_alternatives(db: Session, q_type: str, category: str, difficulty: str,
                      exclude_ids: list[int], limit: int = 30) -> list[Question]:
    """换题用：找同题型（可选同大类/难度）、未被使用的候选题。"""
    pool = [q for q in _apply_filters(db.query(Question), q_type, category, difficulty)
            .order_by(Question.id.desc()).all() if q.id not in set(exclude_ids)]
    return pool[:limit]


# ---------- AB 卷乱序 ----------

def shuffle_for_version(picked: dict[str, list[tuple[Question, int]]], seed: int) -> list[dict]:
    """生成一个卷别的题目编排：节内题序乱序 + 选择题选项乱序。
    返回 [{q_type, label, questions: [{no, stem, options, answer, score, qid}], per_score, total}]"""
    rng = random.Random(seed)
    sections = []
    for q_type in TYPE_ORDER:
        if q_type not in picked:
            continue
        items = list(picked[q_type])
        rng.shuffle(items)
        questions = []
        for no, (q, score) in enumerate(items, start=1):
            stem, options, answer = q.stem, q.options, q.answer
            if q_type in ("single", "multiple") and options:
                # 选项乱序并重排字母，答案随之映射
                opts = list(options)
                rng.shuffle(opts)
                letters = [chr(ord("A") + i) for i in range(len(opts))]
                old_to_new = {o["key"]: letters[i] for i, o in enumerate(opts)}
                options = [{"key": letters[i], "text": o["text"]} for i, o in enumerate(opts)]
                answer = "".join(sorted(old_to_new[ch] for ch in q.answer if ch in old_to_new))
            elif q_type == "judge":
                options = [{"key": "A", "text": "正确"}, {"key": "B", "text": "错误"}]
            questions.append({"no": no, "qid": q.id, "stem": stem, "options": options,
                              "answer": answer, "score": score, "analysis": q.analysis})
        scores = {s for (_, s) in items}
        sections.append({
            "q_type": q_type, "label": TYPE_LABELS[q_type], "questions": questions,
            "per_score": items[0][1] if len(scores) == 1 else None,
            "total": sum(s for (_, s) in items),
        })
    return sections


# ---------- docx 工具 ----------

def _set_cn_font(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _para(doc, text="", size=12, bold=False, align=None, font="宋体", space_after=4):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def _cell_text(cell, text, size=10.5, bold=False, center=True):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "宋体"
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _stem_for_paper(q_type: str, stem: str) -> str:
    stem = stem.rstrip()
    if q_type in ("single", "multiple"):
        if "（" not in stem[-6:]:
            stem += "（    ）"
        return stem
    if q_type == "judge":
        return stem + ("" if stem.endswith(("。", "？", "！")) else "。") + "（    ）"
    if q_type == "fill_blank":
        return stem.replace("_______", "（    ）").replace("______", "（    ）").replace("_____", "（    ）")
    return stem


# ---------- 试卷 docx ----------

def render_paper(title: str, venue: str, version: str, sections: list[dict], path: Path):
    doc = Document()
    _set_cn_font(doc)

    # 封面
    doc.add_paragraph()
    _para(doc, f"《{title}》", size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font="黑体")
    _para(doc, f"试卷（{version}）", size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font="黑体")
    if venue:
        _para(doc, f"（{venue}）", size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    info = doc.add_table(rows=2, cols=4)
    info.style = "Table Grid"
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, label in enumerate(["姓  名", "员工编号", "座位号", "考场号"]):
        _cell_text(info.rows[0].cells[i], label, size=12, bold=True)
        _cell_text(info.rows[1].cells[i], "", size=12)
        info.rows[1].cells[i].height = Cm(1.2)
    doc.add_paragraph()
    _para(doc, "题目必须在答题纸上作答，在题本上作答一律无效",
          size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font="黑体")
    doc.add_page_break()

    # 卷首：标题 + 分数汇总表
    _para(doc, f"{title}试题", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font="黑体")
    summary = doc.add_table(rows=2, cols=len(sections) + 1)
    summary.style = "Table Grid"
    summary.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, sec in enumerate(sections):
        _cell_text(summary.rows[0].cells[i], f"{CN_NUMS[i]}、{sec['label']}", size=10.5, bold=True)
    _cell_text(summary.rows[0].cells[len(sections)], "总分", size=10.5, bold=True)
    doc.add_paragraph()

    # 各节题目
    for i, sec in enumerate(sections):
        n = len(sec["questions"])
        if sec["per_score"] is not None:
            per = (f"每空{sec['per_score']}分" if sec["q_type"] == "fill_blank"
                   else f"每题{sec['per_score']}分")
            head = f"{CN_NUMS[i]}、{sec['label']}（{per}，共{sec['total']}分）"
        else:
            head = f"{CN_NUMS[i]}、{sec['label']}（共{n}题，共{sec['total']}分）"
        _para(doc, head, size=14, bold=True, font="黑体", space_after=6)
        for q in sec["questions"]:
            prefix = (f"{q['no']}.（{q['score']}分）" if sec["per_score"] is None
                      else f"{q['no']}.")
            _para(doc, prefix + _stem_for_paper(sec["q_type"], q["stem"]), space_after=2)
            if sec["q_type"] in ("single", "multiple") and q["options"]:
                for o in q["options"]:
                    _para(doc, f"{o['key']}.{o['text']}", space_after=2)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
    doc.save(path)


# ---------- 答题卡 docx ----------

def _answer_grid(doc, n: int):
    """题号/答案 两行一组、每组最多10题。"""
    for start in range(1, n + 1, 10):
        cols = min(10, n - start + 1)
        t = doc.add_table(rows=2, cols=cols + 1)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        _cell_text(t.rows[0].cells[0], "题号", size=10.5, bold=True)
        _cell_text(t.rows[1].cells[0], "答案", size=10.5, bold=True)
        for j in range(cols):
            _cell_text(t.rows[0].cells[j + 1], str(start + j), size=10.5)
            _cell_text(t.rows[1].cells[j + 1], "", size=10.5)
            t.rows[1].cells[j + 1].height = Cm(0.9)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)


def render_answer_sheet(title: str, venue: str, version: str, sections: list[dict], path: Path):
    doc = Document()
    _set_cn_font(doc)
    _para(doc, f"《{title}》试卷（{version}）", size=15, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, font="黑体")
    _para(doc, "答 题 纸", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font="黑体")
    if venue:
        _para(doc, f"（{venue}）", size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, "请在各组题目答题区域内作答", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    for i, sec in enumerate(sections):
        n = len(sec["questions"])
        per = f"{sec['per_score']}分/题，" if sec["per_score"] is not None else ""
        _para(doc, f"{CN_NUMS[i]}、{sec['label']}（{per}共{n}题，共{sec['total']}分）",
              size=12, bold=True, font="黑体", space_after=4)
        if sec["q_type"] in CHOICE_TYPES:
            _answer_grid(doc, n)
        elif sec["q_type"] == "fill_blank":
            for q in sec["questions"]:
                blanks = max(1, sum(q["stem"].count(u) for u in ("_______", "（    ）")) or
                             len(q["options"] or []) or 1)
                line = f"{q['no']}、" + "  ".join("＿" * 8 for _ in range(blanks))
                _para(doc, line, size=12, space_after=8)
        else:  # 简答/论述：题号+分值+留白
            for q in sec["questions"]:
                _para(doc, f"{q['no']}题（{q['score']}分）", size=12, bold=True, space_after=4)
                for _ in range(6):
                    doc.add_paragraph()
    doc.save(path)


# ---------- 答案与评分参考 docx ----------

def render_answer_key(title: str, versions: dict[str, list[dict]], path: Path):
    doc = Document()
    _set_cn_font(doc)
    _para(doc, f"《{title}》", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font="黑体")
    _para(doc, "参考答案及评分标准（保密）", size=14, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, font="黑体")
    for version, sections in versions.items():
        _para(doc, f"试卷（{version}）", size=14, bold=True, font="黑体", space_after=6)
        for i, sec in enumerate(sections):
            _para(doc, f"{CN_NUMS[i]}、{sec['label']}", size=12, bold=True, font="黑体", space_after=4)
            if sec["q_type"] in CHOICE_TYPES:
                # 紧凑表格：题号/答案
                qs = sec["questions"]
                for start in range(0, len(qs), 10):
                    grp = qs[start:start + 10]
                    t = doc.add_table(rows=2, cols=len(grp) + 1)
                    t.style = "Table Grid"
                    _cell_text(t.rows[0].cells[0], "题号", size=10.5, bold=True)
                    _cell_text(t.rows[1].cells[0], "答案", size=10.5, bold=True)
                    for j, q in enumerate(grp):
                        _cell_text(t.rows[0].cells[j + 1], str(q["no"]), size=10.5)
                        ans = q["answer"]
                        if sec["q_type"] == "judge":
                            ans = {"A": "√", "B": "×"}.get(ans, ans)
                        _cell_text(t.rows[1].cells[j + 1], ans, size=10.5)
                    doc.add_paragraph().paragraph_format.space_after = Pt(2)
            elif sec["q_type"] == "fill_blank":
                for q in sec["questions"]:
                    answers = "；".join(o["text"] for o in (q["options"] or []))
                    _para(doc, f"{q['no']}. {answers}", size=11, space_after=2)
            else:
                for q in sec["questions"]:
                    _para(doc, f"{q['no']}.（{q['score']}分）参考答案：", size=11, bold=True, space_after=2)
                    _para(doc, q["answer"], size=11, space_after=2)
                    if q.get("analysis"):
                        _para(doc, f"评分参考：{q['analysis']}", size=10.5, space_after=4)
        doc.add_page_break()
    doc.save(path)


# ---------- 总入口 ----------

def build_picked_from_ids(db: Session, sections_payload: list[dict]) -> dict[str, list[tuple[Question, int]]]:
    """从前端微调后的显式题目ID构建 picked，用于定稿出卷。
    sections_payload: [{q_type, score, question_ids: [...]}]"""
    picked: dict[str, list[tuple[Question, int]]] = {}
    for sec in sections_payload:
        q_type, score = sec["q_type"], int(sec.get("score", 1))
        ids = sec.get("question_ids", [])
        if not ids:
            continue
        rows = {q.id: q for q in db.query(Question).filter(Question.id.in_(ids)).all()}
        for qid in ids:  # 保持前端给定顺序
            q = rows.get(qid)
            if q and q.status == "approved":
                picked.setdefault(q_type, []).append((q, score))
    if not picked:
        raise ValueError("没有可用题目，请确认所选题目仍为已审核状态")
    return picked


def build_versions(picked: dict[str, list[tuple[Question, int]]],
                   versions: tuple[str, ...]) -> dict[str, list[dict]]:
    """为每个卷别生成乱序后的 sections（A/B 用不同种子，保证乱序不同）。"""
    base_seed = random.randint(0, 10 ** 9)
    seed_offset = {"A": 0, "B": 1, "C": 2, "D": 3}
    return {v: shuffle_for_version(picked, base_seed + seed_offset.get(v, 0)) for v in versions}


def render_picked(title: str, venue: str, picked: dict[str, list[tuple[Question, int]]],
                  versions: tuple[str, ...] = ("A", "B")) -> dict:
    versions = tuple(versions) or ("A",)
    ver_sections = build_versions(picked, versions)

    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    work = EXPORT_DIR / f"paper_{stamp}"
    work.mkdir(parents=True, exist_ok=True)

    files: dict[str, callable] = {}
    for v in versions:
        secs = ver_sections[v]
        files[f"试卷{v}_{title}.docx"] = (lambda p, s=secs, vv=v: render_paper(title, venue, vv, s, p))
        files[f"答题卡{v}_{title}.docx"] = (lambda p, s=secs, vv=v: render_answer_sheet(title, venue, vv, s, p))
    files[f"参考答案_{title}.docx"] = lambda p: render_answer_key(title, ver_sections, p)

    file_names = list(files)
    for fname, fn in files.items():
        fn(work / fname)
    zip_path = EXPORT_DIR / f"组卷_{title}_{stamp}.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for fname in file_names:
            zf.write(work / fname, fname)

    total_q = sum(len(v) for v in picked.values())
    total_score = sum(s for v in picked.values() for (_, s) in v)
    return {"file_path": str(zip_path), "work_dir": str(work), "files": file_names,
            "question_count": total_q, "total_score": total_score,
            "versions": list(versions), "preview": ver_sections}


def assemble_paper(db: Session, title: str, venue: str, criteria: list[dict],
                   versions: tuple[str, ...] = ("A", "B")) -> dict:
    """一步式：按条件抽题并直接出卷（保留给无需微调的快速组卷）。"""
    return render_picked(title, venue, select_questions(db, criteria), versions)
