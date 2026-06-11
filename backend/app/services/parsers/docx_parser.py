from docx import Document as DocxDocument


def parse_docx(file_path: str) -> list[tuple[str, str]]:
    """Word 无可靠页码，按标题分节定位；无标题时按段落块编号。"""
    if file_path.lower().endswith(".doc"):
        raise ValueError("暂不支持旧版 .doc，请先在 Word 中另存为 .docx")
    doc = DocxDocument(file_path)
    segments: list[tuple[str, str]] = []
    current_title = ""
    buf: list[str] = []
    block_no = 0

    def flush():
        nonlocal block_no
        if buf:
            block_no += 1
            locator = current_title or f"段落块{block_no}"
            segments.append((locator, "\n".join(buf)))
            buf.clear()

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name.startswith(("Heading", "标题")):
            flush()
            current_title = text[:60]
            buf.append(text)
        else:
            buf.append(text)
    flush()

    # 表格内容附在末尾
    for ti, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            segments.append((f"表格{ti}", "\n".join(rows)))
    return segments
