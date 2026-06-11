# 端到端冒烟测试：登录→上传→解析→生成(mock)→审核→导出
import io
import sys
import time

import httpx
from docx import Document

sys.stdout.reconfigure(encoding="utf-8")
BASE = "http://127.0.0.1:8000/api/v1"


def make_test_docx() -> bytes:
    doc = Document()
    doc.add_heading("数智化转型培训资料", level=1)
    doc.add_paragraph("中国移动坚持实施网络强基、全栈创新战略，推进AI First、Cloud Fast、Mobile Most布局。")
    doc.add_heading("第一章 AI Agent 应用场景", level=2)
    doc.add_paragraph("在构建场景化的AI Agent时，业务场景应具备高重复性、依赖经验的持续沉淀与快速优化、"
                      "传统路径覆盖模式难以穷举与扩展等特点。告警分析环节中，使用大模型进行告警级别分类"
                      "和低级别告警剔除，可以提升网络运维效率。")
    doc.add_heading("第二章 通信服务构成", level=2)
    doc.add_paragraph("通信服务主要由移动通信、宽带网络、蜂窝物联、卫星互联构成。"
                      "战略价值包括客户满意度、品牌认知、质量提升等多个维度。")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def main():
    c = httpx.Client(timeout=60)
    # 1. 登录
    r = c.post(f"{BASE}/auth/login", json={"username": "admin", "password": "admin123"})
    assert r.status_code == 200, r.text
    c.headers["Authorization"] = f"Bearer {r.json()['token']}"
    print("1. 登录 OK")

    # 2. 上传
    r = c.post(f"{BASE}/documents", files={"file": ("测试培训资料.docx", make_test_docx(),
               "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
    assert r.status_code == 200, r.text
    doc_id = r.json()["id"]
    print(f"2. 上传 OK doc_id={doc_id}")

    # 3. 等待解析
    for _ in range(20):
        time.sleep(0.5)
        d = c.get(f"{BASE}/documents/{doc_id}").json()
        if d["parse_status"] in ("done", "failed"):
            break
    assert d["parse_status"] == "done", d
    print(f"3. 解析 OK，切片数={d['chunk_count']}")
    chunks = c.get(f"{BASE}/documents/{doc_id}/chunks").json()
    for ch in chunks:
        print(f"   切片[{ch['source_locator']}] {ch['content'][:30]}...")

    # 4. 创建生成任务（mock 模型）
    r = c.post(f"{BASE}/tasks", json={"document_id": doc_id,
               "type_counts": {"single": 2, "multiple": 1, "judge": 2, "short_answer": 1},
               "difficulty": "medium"})
    assert r.status_code == 200, r.text
    task_id = r.json()["id"]
    for _ in range(30):
        time.sleep(0.5)
        t = c.get(f"{BASE}/tasks/{task_id}").json()
        if t["status"] in ("done", "failed"):
            break
    assert t["status"] == "done", t
    print(f"4. 生成 OK，题目数={t['question_count']}，模型={t['model_name']}")

    # 5. 审核：列待审 → 编辑一题 → 通过全部
    qs = c.get(f"{BASE}/questions", params={"status": "pending_review"}).json()
    assert qs["total"] >= 5, qs["total"]
    first = qs["items"][0]
    detail = c.get(f"{BASE}/questions/{first['id']}").json()
    assert detail["source"] is not None, "题目应能溯源到切片"
    print(f"5a. 溯源 OK：题目{first['id']} ← {detail['source']['source_locator']}")
    r = c.put(f"{BASE}/questions/{first['id']}", json={"stem": first["stem"] + "（已人工修订）"})
    assert r.status_code == 200, r.text
    ids = [q["id"] for q in qs["items"]]
    r = c.post(f"{BASE}/questions/batch-review", json={"ids": ids, "action": "approve"})
    assert r.json()["count"] == len(ids), r.text
    logs = c.get(f"{BASE}/questions/{first['id']}/logs").json()
    assert any(l["action"] == "edit" for l in logs) and any(l["action"] == "approve" for l in logs)
    print(f"5b. 审核 OK：通过 {len(ids)} 题，审核日志 {len(logs)} 条")

    # 6. 导出 Excel 并校验格式
    r = c.post(f"{BASE}/exports", json={})
    assert r.status_code == 200, r.text
    exp = r.json()
    r = c.get(f"{BASE}{exp['download_url'].removeprefix('/api/v1')}")
    assert r.status_code == 200
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(r.content))
    ws = wb["试题导入模板"]
    headers = [cell.value for cell in ws[1]]
    assert headers == ["题干", "类型", "答案", "选项A", "选项B", "选项C", "选项D"], headers
    types = {row[1].value for row in ws.iter_rows(min_row=2)}
    print(f"6. 导出 OK：{exp['question_count']} 题，题型={types}")
    for row in ws.iter_rows(min_row=2, max_row=4, values_only=True):
        print("   ", row)

    print("\n=== 全流程端到端测试通过 ===")


if __name__ == "__main__":
    main()
