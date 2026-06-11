import io
import sys
import zipfile

import httpx
from docx import Document

sys.stdout.reconfigure(encoding="utf-8")
BASE = "http://127.0.0.1:8000/api/v1"
c = httpx.Client(timeout=60)
r = c.post(f"{BASE}/auth/login", json={"username": "admin", "password": "admin123"})
c.headers["Authorization"] = f"Bearer {r.json()['token']}"
papers = c.get(f"{BASE}/papers").json()
r = c.get(f"{BASE}{papers[0]['download_url'].removeprefix('/api/v1')}")
zf = zipfile.ZipFile(io.BytesIO(r.content))
name = [n for n in zf.namelist() if n.startswith("参考答案")][0]
doc = Document(io.BytesIO(zf.read(name)))
print("=== paragraphs ===")
for p in doc.paragraphs[:30]:
    print(repr(p.text))
print("=== tables ===")
for t in doc.tables:
    for row in t.rows:
        print("|".join(c.text for c in row.cells))
    print("---")
